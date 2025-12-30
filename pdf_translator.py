import os
import sys
import argparse
import requests
import time
import tempfile
from pdf2docx import Converter
from docx import Document
from docx2pdf import convert

def translate_text_with_deepl(text, target_lang, auth_key):
    """Translate text using DeepL API."""
    if not text or text.strip() == '':
        return ''

    url = "https://api-free.deepl.com/v2/translate"

    # If you have a DeepL Pro subscription, use this URL instead:
    # url = "https://api.deepl.com/v2/translate"

    # DeepL has character limits per request, so we might need to split large texts
    max_chars = 5000
    if len(text) <= max_chars:
        chunks = [text]
    else:
        # Split by sentences to avoid cutting in the middle of a sentence
        sentences = text.replace('. ', '.|').replace('! ', '!|').replace('? ', '?|').split('|')
        chunks = []
        current_chunk = ""

        for sentence in sentences:
            if len(current_chunk) + len(sentence) + 1 <= max_chars:
                current_chunk += sentence + " "
            else:
                chunks.append(current_chunk.strip())
                current_chunk = sentence + " "

        if current_chunk:
            chunks.append(current_chunk.strip())

    translated_text = ""

    for chunk in chunks:
        payload = {
            "text": [chunk],
            "target_lang": target_lang
        }

        headers = {
            "Authorization": f"DeepL-Auth-Key {auth_key}",
            "Content-Type": "application/json"
        }

        try:
            response = requests.post(url, json=payload, headers=headers)
            response.raise_for_status()
            result = response.json()
            translated_chunk = result["translations"][0]["text"]
            translated_text += translated_chunk + " "

            # Respect DeepL API rate limits
            time.sleep(0.5)

        except requests.exceptions.RequestException as e:
            print(f"Translation error: {e}")
            if hasattr(response, 'status_code'):
                if response.status_code == 429:  # Too Many Requests
                    print("Rate limit exceeded. Waiting before retrying...")
                    time.sleep(60)  # Wait a minute before retrying
                elif response.status_code == 456:  # Quota exceeded
                    print("DeepL API quota exceeded.")
                    raise
            else:
                raise

    return translated_text.strip()

def translate_pdf_via_docx(pdf_path, target_lang, auth_key, output_path=None):
    """
    Translate PDF by converting to DOCX, translating text, then converting back to PDF.
    This preserves images and formatting.
    """
    # Set default output path if not specified
    if not output_path:
        filename = os.path.basename(pdf_path)
        base_name = os.path.splitext(filename)[0]
        output_path = f"{base_name}_{target_lang}.pdf"

    # Create temporary directory for intermediate files
    # with tempfile.TemporaryDirectory() as temp_dir:  ##########
    temp_dir = tempfile.mkdtemp(dir=".")  ##########

    # Temporary file paths
    temp_docx = os.path.join(temp_dir, "temp_document.docx")
    translated_docx = os.path.join(temp_dir, "translated_document.docx")

    print(f"Converting PDF to DOCX: {pdf_path} -> {temp_docx}")
    # Convert PDF to DOCX
    cv = Converter(pdf_path)
    cv.convert(temp_docx)
    cv.close()

    print("DOCX conversion complete. Starting translation...")
    # Open the DOCX file
    doc = Document(temp_docx)

    # Translate paragraphs
    total_paragraphs = len(doc.paragraphs)
    print(f"Total paragraphs to translate: {total_paragraphs}")
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            print(f"Translating paragraph {i+1}/{total_paragraphs}...")
            translated_text = translate_text_with_deepl(paragraph.text, target_lang, auth_key)
            paragraph.text = translated_text

    # Translate tables
    tables_count = len(doc.tables)
    print(f"Total tables to translate: {tables_count}")
    for table_idx, table in enumerate(doc.tables):
        print(f"Translating table {table_idx+1}/{tables_count}...")
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        translated_text = translate_text_with_deepl(paragraph.text, target_lang, auth_key)
                        paragraph.text = translated_text

    # Translate text boxes and other elements in headers/footers
    sections_count = len(doc.sections)
    print(f"Total sections to check for headers/footers: {sections_count}")
    for section_idx, section in enumerate(doc.sections):
        print(f"Translating headers and footers in section {section_idx+1}/{sections_count}...")

        # Process header
        for paragraph in section.header.paragraphs:
            if paragraph.text.strip():
                translated_text = translate_text_with_deepl(paragraph.text, target_lang, auth_key)
                paragraph.text = translated_text

        # Process footer
        for paragraph in section.footer.paragraphs:
            if paragraph.text.strip():
                translated_text = translate_text_with_deepl(paragraph.text, target_lang, auth_key)
                paragraph.text = translated_text

    # Save the translated DOCX
    print(f"Saving translated DOCX: {translated_docx}")
    doc.save(translated_docx)

    # Convert back to PDF
    print(f"Converting translated DOCX to PDF: {translated_docx} -> {output_path}")
    try:
        convert(translated_docx, output_path)
        print(f"Translation completed! Output saved to: {output_path}")
    except Exception as e:
        print(f"Error converting DOCX to PDF: {e}")
        print(f"Saving translated DOCX only to: {output_path.replace('.pdf', '.docx')}")
        # Save as DOCX if PDF conversion fails
        import shutil
        shutil.copy2(translated_docx, output_path.replace('.pdf', '.docx'))

def main():
    parser = argparse.ArgumentParser(description='Translate PDF files using DeepL API while preserving formatting and images')
    parser.add_argument('pdf_path', help='Path to the PDF file to translate')
    parser.add_argument('target_lang', help='Target language code (e.g., EN, DE, FR, etc.)')
    parser.add_argument('--auth-key', help='DeepL API authentication key', required=True)
    parser.add_argument('--output', help='Path for the translated PDF', default=None)

    args = parser.parse_args()

    # Check if the input PDF exists
    if not os.path.isfile(args.pdf_path):
        print(f"Error: The input PDF file '{args.pdf_path}' does not exist.")
        sys.exit(1)

    translate_pdf_via_docx(args.pdf_path, args.target_lang, args.auth_key, args.output)

if __name__ == "__main__":
    main()


# Example usage:
# python pdf_translator.py epyramid_infant_toddler_leader_guide_5.pdf ES --auth-key XXX --output epyramid_infant_toddler_leader_guide_5_es.pdf